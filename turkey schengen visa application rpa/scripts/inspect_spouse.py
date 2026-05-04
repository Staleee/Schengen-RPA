"""Print every paragraph + table cell that contains the word 'Spouse' or
'Children' (sections 35 and 36) in visaform.docx, along with neighbours and any
single-brace placeholders inside that area. Used to design a patch."""

from pathlib import Path

from docx import Document

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "visaform.docx"


def cell_text(cell) -> str:
    return "".join(p.text for p in cell.paragraphs)


def main() -> None:
    doc = Document(str(TEMPLATE_PATH))

    print("=== Body paragraphs near 'Spouse' / 'Children' ===")
    paragraphs = list(doc.paragraphs)
    for i, p in enumerate(paragraphs):
        text = "".join(r.text for r in p.runs)
        if "Spouse" in text or "Children" in text or "35." in text or "36." in text:
            for j in range(max(0, i - 1), min(len(paragraphs), i + 3)):
                pj_text = "".join(r.text for r in paragraphs[j].runs)
                print(f"  body[{j}] {repr(pj_text[:120])}")
            print()

    print("=== Tables containing 'Spouse' / 'Children' ===")
    for ti, table in enumerate(doc.tables):
        body = [cell_text(c) for row in table.rows for c in row.cells]
        if any("Spouse" in t or "Children" in t or "35." in t or "36." in t for t in body):
            print(f"--- Table {ti} ({len(table.rows)} rows) ---")
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    txt = cell_text(cell)
                    if not txt.strip():
                        continue
                    print(f"  T{ti}R{ri}C{ci}: {repr(txt[:120])}")


if __name__ == "__main__":
    main()
