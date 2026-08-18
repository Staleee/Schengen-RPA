"""
One-time patch: make the "22. Number of entries requested" checkboxes data-driven.

The template hard-coded a checked box before "Single Entry" ({cbc}) and an empty box before
"Multiple entry" ({cbe}), so choosing Multiple still showed Single checked. This replaces them
with the dynamic {entry_single} / {entry_multiple} placeholders that turkey_docx_fill fills
from `number_of_entries`.

Idempotent: skips when the placeholders are already present.

Run from the turkey RPA directory:
    python scripts/patch_turkey_entries.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

PACKAGE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PACKAGE_DIR / "visaform.docx"


def _set_text(para, new_text: str) -> None:
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _patch_paragraph(para) -> bool:
    """Replace the hard-coded entry checkboxes in the §22 line. Returns True if changed."""
    full = "".join(r.text for r in para.runs)
    if "{entry_single}" in full or "{entry_multiple}" in full:
        return False  # already patched
    # Only touch the entries line that has a checked box before "Single Entry".
    if "{cbc}" not in full or "Single Entry" not in full:
        return False
    new = full.replace("{cbc}", "{entry_single}").replace("{cbe}", "{entry_multiple}")
    _set_text(para, new)
    return True


def patch(path: Path) -> None:
    if not path.exists():
        print(f"  ERROR: template not found at {path}")
        return
    doc = Document(str(path))
    changed = 0

    def walk(paragraphs):
        nonlocal changed
        for p in paragraphs:
            if _patch_paragraph(p):
                changed += 1

    walk(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                walk(cell.paragraphs)
    for section in doc.sections:
        for block in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            if block is not None:
                walk(block.paragraphs)

    if changed:
        doc.save(str(path))
        print(f"  PATCHED {changed} paragraph(s) in {path.name}")
    else:
        print(f"  OK (already patched or no matching line): {path.name}")


if __name__ == "__main__":
    print("Patching Turkey template entry checkboxes...")
    patch(TEMPLATE_PATH)
    print("Done.")
