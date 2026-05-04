"""
Diagnostic + patch script for visaform.docx.

1. Lists all {placeholder} tokens found in the template.
2. Checks whether {rvisa_number}, {current_citizenship}, {citizenship_at_birth} are present.
3. If any are missing, scans table cells for the nearest label (e.g. "residence", "visa no",
   "citizenship", "nationality") and inserts the placeholder into the adjacent value cell.
4. Saves the modified template (idempotent: skips cells that already have the placeholder).

Run once locally or during Docker build (python-docx required):
    python scripts/patch_visaform_placeholders.py
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PACKAGE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PACKAGE_DIR / "visaform.docx"

_SINGLE_BRACE_RE = re.compile(r"\{([^{}]+)\}")

# Mapping: placeholder we need -> keywords (lowercase) to find in an adjacent label cell.
REQUIRED_PLACEHOLDERS: dict[str, list[str]] = {
    "rvisa_number": ["rvisa", "r visa", "residence permit", "visa no", "visa number", "uae visa"],
    "rvisa_expiry_date": ["rvisa expiry", "r visa expiry", "residence permit expiry", "visa expiry"],
    "current_citizenship": ["current citizenship", "citizenship", "nationality"],
    "citizenship_at_birth": ["citizenship at birth", "birth citizenship", "original citizenship"],
}


def _cell_text(cell) -> str:
    return "".join(p.text for p in cell.paragraphs).strip()


def _find_placeholders(doc: Document) -> set[str]:
    found = set()
    all_text = []
    for para in doc.paragraphs:
        all_text.append("".join(r.text for r in para.runs))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text.append("".join(r.text for r in para.runs))
    for text in all_text:
        for m in _SINGLE_BRACE_RE.finditer(text):
            found.add(m.group(1).strip().lower().replace(" ", "_").replace("-", "_"))
    return found


def _label_matches(label: str, keywords: list[str]) -> bool:
    label_lower = label.lower()
    return any(kw in label_lower for kw in keywords)


def _append_placeholder_to_cell(cell, placeholder: str) -> None:
    """Append {placeholder} text to the last paragraph of a cell."""
    # Use the last paragraph (or add one)
    if cell.paragraphs:
        para = cell.paragraphs[-1]
    else:
        para = cell.add_paragraph()
    run = para.add_run(f"{{{placeholder}}}")
    # Don't apply bold – let fill_turkey_docx_bytes handle it naturally


def _row_has_any_placeholder(row, exclude: set[str] | None = None) -> bool:
    """Return True if any cell in this row contains a `{...}` placeholder.

    `exclude` is a set of placeholder *names* (without braces) to ignore — used so
    decorative checkbox markers like `{cbe}`/`{cbc}` don't count as a real value
    for the labeled field.
    """
    exclude = exclude or set()
    for cell in row.cells:
        for para in cell.paragraphs:
            text = "".join(r.text for r in para.runs)
            for m in _SINGLE_BRACE_RE.finditer(text):
                name = m.group(1).strip()
                if name not in exclude:
                    return True
    return False


_DECORATIVE_PLACEHOLDERS = {"cbe", "cbc", "cbo"}


def _try_add_in_tables(doc: Document, placeholder: str, keywords: list[str]) -> bool:
    """
    Try to find a table row where one cell is a label matching keywords,
    and the *same row* has no real placeholder yet.
    Returns True if the placeholder was inserted (or is already covered).

    IMPORTANT: We deliberately scope the check to the same row as the matching
    label. The previous behaviour walked sibling cells across the row and would
    happily plant `{current_citizenship}` inside the Spouse section's
    "Nationality" cell because the keyword list included "nationality". Now we:

    - First, look for a value cell within the same row that already contains a
      real (non-decorative) placeholder. If one exists, treat the field as
      already covered and DO NOT insert anything.
    - Otherwise, append the placeholder to the cell directly to the right of the
      label cell. We never wander into other rows.
    """
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            label_cell_index: int | None = None
            for i, cell in enumerate(cells):
                if _label_matches(_cell_text(cell), keywords):
                    label_cell_index = i
                    break
            if label_cell_index is None:
                continue

            # Already covered (by an existing placeholder in the row)?
            if _row_has_any_placeholder(row, exclude=_DECORATIVE_PLACEHOLDERS):
                # Verify the existing placeholder isn't just decorative — if so
                # the labeled field is filled by some other key in
                # build_replacements and we should leave the template alone.
                return True

            # Insert into the cell immediately after the label.
            target_index = label_cell_index + 1
            if target_index >= len(cells):
                continue
            val_cell = cells[target_index]
            _append_placeholder_to_cell(val_cell, placeholder)
            label = _cell_text(cells[label_cell_index])
            print(f"    Inserted {{{placeholder}}} in cell after '{label}'")
            return True
    return False


# `build_replacements` substitutes `{visa_duration}` with values like "15 days" (it
# always appends the unit). Templates that have a hardcoded " days" word right after
# the placeholder render as "15 days days". Strip those duplicate units.
_DUPLICATE_DAYS_RE = re.compile(r"(\{visa_duration\})\s+days\b", re.IGNORECASE)

def _strip_duplicate_days(doc: Document) -> int:
    """Find paragraphs (incl. table cells) where `{visa_duration}` is followed by a
    literal ' days' and remove the redundant word. Returns the number of edits."""
    edits = 0

    def fix_paragraph(paragraph) -> None:
        nonlocal edits
        full_text = "".join(run.text for run in paragraph.runs)
        if "{visa_duration}" not in full_text.lower():
            return
        new_text = _DUPLICATE_DAYS_RE.sub(r"\1", full_text)
        if new_text == full_text:
            return
        for run in list(paragraph.runs):
            run._r.getparent().remove(run._r)
        paragraph.add_run(new_text)
        edits += 1

    for para in doc.paragraphs:
        fix_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_paragraph(para)
    return edits


def patch_template(path: Path) -> None:
    if not path.exists():
        print(f"SKIP (not found): {path}")
        sys.exit(1)

    doc = Document(path)
    existing = _find_placeholders(doc)

    print(f"\nTemplate: {path.name}")
    print(f"Existing placeholders ({len(existing)}): {sorted(existing)}")

    changed = False
    for placeholder, keywords in REQUIRED_PLACEHOLDERS.items():
        if placeholder in existing:
            print(f"  OK: {{{placeholder}}}")
        else:
            print(f"  MISSING: {{{placeholder}}} \u2014 searching for label matching {keywords[:2]}...")
            inserted = _try_add_in_tables(doc, placeholder, keywords)
            if inserted:
                changed = True
            else:
                print(f"    WARNING: Could not find a suitable cell. Add {{{placeholder}}} manually.")

    duplicate_edits = _strip_duplicate_days(doc)
    if duplicate_edits:
        print(f"  Removed redundant ' days' word after {{visa_duration}} ({duplicate_edits} edit(s))")
        changed = True

    if changed:
        doc.save(str(path))
        print(f"Saved patched template: {path.name}")
    else:
        print("No changes needed.")


if __name__ == "__main__":
    patch_template(TEMPLATE_PATH)
