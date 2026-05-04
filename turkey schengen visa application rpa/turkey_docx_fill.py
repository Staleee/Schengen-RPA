"""
Turkey Schengen – fill a Word template using single-brace placeholders {name}.

Checkbox sections (6 Sex, 8 Marital status): put tokens **in the box or just before the label**
(same line is fine). The API fills them from `sex` and `marital_status` in the JSON.

**Short (minimal text in Word):**
  {6m} {6f}  — male / female
  {8s} {8m}  — single / married (mnemonic letters)
  {8a} {8b} {8c} {8d} {8e} {8f}  — same order: single → other (if you prefer letters a–f)

**Long (same behaviour):**
  {sex_check_male} {sex_check_female}
  {marital_check_single} … {marital_check_other}

**§24 / §25 (Turkey travel history):** `{24y}` `{24n}` — visited Turkey before yes/no; `{25y}` `{25n}` — deported/refused from Turkey yes/no.
JSON (preferred): `maid_traveled_to_turkey_before`, `maid_deported_from_turkey_before` (`yes`/`no` or bool). Legacy: `traveled_turkey_before`, `deported_from_turkey_before`.

**Client email typo in Word:** `{client)email}` maps to the same value as `client_email`.

**Visa duration:** When `arrival_date` and `departure_date` parse, **`visa_duration` is always set from them** (inclusive day count) and **overrides** any value sent in the JSON.

**Means of transport:** Send `means_of_transport` (e.g. `Air`) for `{means_of_transport}` in Word.

**`{maid_country_of_birth}`:** Always **derived** from `maid_nationality` (and a few alternate keys), not taken from the request body for this field — e.g. Filipino → Philippines. Unknown demonyms yield an empty string.

**Static checkboxes (passport type, etc.):** No JSON — use `{cbe}` (**e**mpty) or `{cbc}` (**c**hecked). Same meaning: `{checkbox_empty}` / `{checkbox_checked}`, `{cb_empty}` / `{cb_checked}`.

**Dynamic checkboxes:** `{6m}` `{6f}`; `{8s}` `{8m}` or `{8a}`…`{8f}`; `{sex_check_*}`; `{marital_check_*}`; `{24y}` `{24n}` `{25y}` `{25n}` — from `sex`, `marital_status`, Turkey history. All use Unicode **☐** / **☑** in text runs (same as before Webdings experiments).

Optional: **TURKEY_CHECKBOX_FONT** forces the font name for runs that contain ☐/☑ (default Segoe UI Symbol on Windows, DejaVu Sans elsewhere).
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

_INVALID_XML_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\ufffe\uffff]")
_SINGLE_BRACE_RE = re.compile(r"\{([^{}]+)\}")

ASSETS_DIR = Path(__file__).resolve().parent / "assets"
PACKAGE_DIR = Path(__file__).resolve().parent


def normalize_key(text: str) -> str:
    if not text or not isinstance(text, str):
        return ""
    t = text.strip()
    t = re.sub(r"([a-z0-9])([A-Z])", r"\1_\2", t)
    t = re.sub(r"([A-Z]+)([A-Z][a-z])", r"\1_\2", t)
    t = t.lower()
    t = re.sub(r"\s+", "_", t)
    t = re.sub(r"[^\w\-]", "", t)
    return t


def _sanitize_for_word(value: str) -> str:
    if not value:
        return ""
    return _INVALID_XML_RE.sub("", value)


def resolve_word_template() -> Path:
    env = os.environ.get("TURKEY_WORD_TEMPLATE")
    if env:
        p = Path(env).expanduser()
        if p.is_file():
            return p
    preferred_names = ("tvisaform.docx", "visaform.docx", "turkey_application.docx")
    search_roots = (ASSETS_DIR, PACKAGE_DIR)
    for root in search_roots:
        for name in preferred_names:
            p = root / name
            if p.is_file():
                return p
    for root in search_roots:
        for p in sorted(root.glob("*.docx")):
            if p.name.startswith("~$"):
                continue
            return p
    return ASSETS_DIR / "tvisaform.docx"


def _normalize_marital(raw: str) -> str:
    s = (raw or "").strip().lower().replace(" ", "_")
    aliases = {
        "divorce": "divorced",
        "widow": "widowed",
        "widower": "widowed",
        "spouse": "married",
    }
    return aliases.get(s, s)


_CHECK_ON = "☑"
_CHECK_OFF = "☐"
_CHECK_GLYPHS = frozenset((_CHECK_ON, _CHECK_OFF))


def _reserved_checkbox_glyph_placeholders() -> Dict[str, str]:
    """Always available in Word: static empty/checked box (not driven by JSON)."""
    return {
        "cbe": _CHECK_OFF,
        "cbc": _CHECK_ON,
        "checkbox_empty": _CHECK_OFF,
        "checkbox_checked": _CHECK_ON,
        "cb_empty": _CHECK_OFF,
        "cb_checked": _CHECK_ON,
    }


def _font_name_for_run_text(text: str, base_font: Optional[str]) -> Optional[str]:
    """Runs containing ☐/☑ use a symbol font so glyphs are not paper-thin in some body fonts."""
    if text and any(c in _CHECK_GLYPHS for c in text):
        return (os.environ.get("TURKEY_CHECKBOX_FONT") or "").strip() or (
            "Segoe UI Symbol" if os.name == "nt" else "DejaVu Sans"
        )
    return base_font


def checkbox_placeholders(sex: Any, marital_status: Any) -> Dict[str, str]:
    """Checked / unchecked Unicode boxes for Word (sections 6 and 8)."""
    uchk, chk = _CHECK_OFF, _CHECK_ON
    s = str(sex or "").strip().lower()
    if s in ("m", "male", "man", "1"):
        male, female = chk, uchk
    elif s in ("f", "female", "woman", "2"):
        male, female = uchk, chk
    else:
        male, female = uchk, uchk

    m = _normalize_marital(str(marital_status or ""))
    opts = ("single", "married", "separated", "divorced", "widowed", "other")
    out = {
        "sex_check_male": male,
        "sex_check_female": female,
    }
    for opt in opts:
        out[f"marital_check_{opt}"] = chk if m == opt else uchk
    # Compact §6 / §8 tokens (same values as long names; easier on layout)
    out["6m"] = out["sex_check_male"]
    out["6f"] = out["sex_check_female"]
    short_marital = ("8a", "8b", "8c", "8d", "8e", "8f")
    for letter, opt in zip(short_marital, opts):
        out[letter] = out[f"marital_check_{opt}"]
    # Mnemonic single / married (common on short forms)
    out["8s"] = out["marital_check_single"]
    out["8m"] = out["marital_check_married"]
    return out


def _truthy_tristate(v: Any) -> Optional[bool]:
    """True/False or None if unknown / empty."""
    if v is None:
        return None
    if isinstance(v, bool):
        return v
    s = str(v).strip().lower()
    if s in ("", "unknown", "n/a", "na", "null"):
        return None
    if s in ("y", "yes", "true", "1", "on"):
        return True
    if s in ("n", "no", "false", "0", "off"):
        return False
    return None


def turkey_history_checkboxes(flat: Dict[str, Any]) -> Dict[str, str]:
    """§24 visited Turkey before → {24y}/{24n}; §25 deported/refused → {25y}/{25n}."""
    uchk, chk = _CHECK_OFF, _CHECK_ON
    visited = _truthy_tristate(
        flat.get("maid_traveled_to_turkey_before")
        or flat.get("traveled_turkey_before")
        or flat.get("been_to_turkey_before")
        or flat.get("turkey_visited_before")
    )
    if visited is True:
        y24, n24 = chk, uchk
    elif visited is False:
        y24, n24 = uchk, chk
    else:
        y24, n24 = uchk, uchk

    deported = _truthy_tristate(
        flat.get("maid_deported_from_turkey_before")
        or flat.get("maid_deported_from_tueky_before")
        or flat.get("deported_from_turkey_before")
        or flat.get("deported_from_turkey")
        or flat.get("turkey_deported_before")
    )
    if deported is True:
        y25, n25 = chk, uchk
    elif deported is False:
        y25, n25 = uchk, chk
    else:
        y25, n25 = uchk, uchk

    return {"24y": y24, "24n": n24, "25y": y25, "25n": n25}


# Zoho often sends dd.MM.yyyy — try dotted forms first to avoid any ambiguity with slashes.
_DATE_FORMATS = (
    "%d.%m.%Y",
    "%d.%m.%y",
    "%d/%m/%Y",
    "%d/%m/%y",
    "%Y-%m-%d",
    "%d-%m-%Y",
    "%d %B %Y",
    "%d %b %Y",
)


def _normalize_date_string(s: str) -> str:
    """Make Zoho / copy-paste variants parseable (Unicode dots, thin spaces)."""
    t = s.replace("\u00a0", " ").replace("\u2009", " ").strip()
    t = t.replace("．", ".").replace("·", ".").replace("∙", ".")
    # "01,04,2026" or "01-04-2026" already handled by format list; optional comma as sep
    if re.match(r"^\d{1,2},\d{1,2},\d{4}$", t):
        t = t.replace(",", ".")
    return t


def _parse_flexible_date(val: Any) -> Optional[date]:
    if val is None:
        return None
    if isinstance(val, date) and not isinstance(val, datetime):
        return val
    if isinstance(val, datetime):
        return val.date()
    s = str(val).strip()
    if not s:
        return None
    # ISO date or datetime prefix: 2026-04-01 or 2026-04-01T12:00:00Z
    if len(s) >= 10 and s[4:5] == "-" and s[7:8] == "-":
        try:
            return datetime.strptime(s[:10], "%Y-%m-%d").date()
        except ValueError:
            pass
    s = _normalize_date_string(s)
    # Zoho year/month/day: 2026/4/15 or 2026-04-15 (avoid misreading as dd/mm)
    m = re.fullmatch(r"(\d{4})[/.](\d{1,2})[/.](\d{1,2})", s)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            return date(y, mo, d)
        except ValueError:
            pass
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


# Start of stay in destination: NOT start_date (Zoho noise). See compute_inclusive_stay_days.
_VISA_ARRIVAL_KEYS_ORDER: Tuple[str, ...] = (
    "arrival_date",
    "arrivaldate",
    "date_of_arrival",
    "travel_arrival_date",
)
# End of stay when arrival_* is already set (leave Schengen / Turkey or return flight).
_VISA_END_AFTER_ARRIVAL_ORDER: Tuple[str, ...] = (
    "departure_date",
    "departuredate",
    "date_of_departure",
    "travel_departure_date",
    "return_date",
    "returndate",
    "date_of_return",
)


def _flat_first_by_normalized_key(flat: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for k, v in flat.items():
        nk = normalize_key(str(k))
        if nk and nk not in out:
            out[nk] = v
    return out


# Normalized nationality / country wording (lowercase) → country of birth for `{maid_country_of_birth}`.
_NATIONALITY_TO_COUNTRY_OF_BIRTH: Dict[str, str] = {
    "filipino": "Philippines",
    "filipina": "Philippines",
    "philippine": "Philippines",
    "philippines": "Philippines",
    "phillipines": "Philippines",
    "phillippines": "Philippines",
    "indian": "India",
    "india": "India",
    "nepali": "Nepal",
    "nepalese": "Nepal",
    "nepal": "Nepal",
    "sri lankan": "Sri Lanka",
    "srilankan": "Sri Lanka",
    "bangladeshi": "Bangladesh",
    "bangladesh": "Bangladesh",
    "pakistani": "Pakistan",
    "pakistan": "Pakistan",
    "indonesian": "Indonesia",
    "indonesia": "Indonesia",
    "ethiopian": "Ethiopia",
    "ethiopia": "Ethiopia",
    "ugandan": "Uganda",
    "uganda": "Uganda",
    "kenyan": "Kenya",
    "kenya": "Kenya",
    "egyptian": "Egypt",
    "egypt": "Egypt",
    "sudanese": "Sudan",
    "sudan": "Sudan",
    "syrian": "Syria",
    "syria": "Syria",
    "jordanian": "Jordan",
    "jordan": "Jordan",
    "lebanese": "Lebanon",
    "lebanon": "Lebanon",
    "vietnamese": "Vietnam",
    "vietnam": "Vietnam",
    "viet nam": "Vietnam",
    "chinese": "China",
    "china": "China",
    "thai": "Thailand",
    "thailand": "Thailand",
    "burmese": "Myanmar",
    "myanmar": "Myanmar",
    "cambodian": "Cambodia",
    "cambodia": "Cambodia",
    "laotian": "Laos",
    "lao": "Laos",
    "laos": "Laos",
    "malaysian": "Malaysia",
    "malaysia": "Malaysia",
    "nigerian": "Nigeria",
    "nigeria": "Nigeria",
    "ghanaian": "Ghana",
    "ghana": "Ghana",
    "turkish": "Turkey",
    "turkey": "Turkey",
    "turkmen": "Turkmenistan",
    "turkmenistan": "Turkmenistan",
    "iranian": "Iran",
    "iran": "Iran",
    "iraqi": "Iraq",
    "iraq": "Iraq",
    "afghan": "Afghanistan",
    "afghanistani": "Afghanistan",
    "afghanistan": "Afghanistan",
    "russian": "Russia",
    "russia": "Russia",
    "ukrainian": "Ukraine",
    "ukraine": "Ukraine",
    "british": "United Kingdom",
    "english": "United Kingdom",
    "uk": "United Kingdom",
    "united kingdom": "United Kingdom",
    "american": "United States",
    "u.s.": "United States",
    "us": "United States",
    "usa": "United States",
    "united states": "United States",
}

_NATIONALITY_SOURCE_KEYS: Tuple[str, ...] = (
    "maid_nationality",
    "nationality",
    "maid_nationality_in_passport",
    "nationality_of_maid",
    "worker_nationality",
)


def _nationality_raw_from_flat(flat: Dict[str, Any]) -> str:
    by_nk = _flat_first_by_normalized_key(flat)
    for nk in _NATIONALITY_SOURCE_KEYS:
        if nk not in by_nk:
            continue
        v = by_nk[nk]
        if v is None:
            continue
        s = str(v).strip()
        if s:
            return s
    return ""


def maid_country_of_birth_from_nationality(flat: Dict[str, Any]) -> str:
    """
    Country name for `{maid_country_of_birth}` from nationality text (not from a dedicated body field).
    """
    raw = _nationality_raw_from_flat(flat)
    if not raw:
        return ""
    key = re.sub(r"\s+", " ", raw.strip().lower())
    if key in _NATIONALITY_TO_COUNTRY_OF_BIRTH:
        return _NATIONALITY_TO_COUNTRY_OF_BIRTH[key]
    if key.startswith("the "):
        key = key[4:].strip()
        if key in _NATIONALITY_TO_COUNTRY_OF_BIRTH:
            return _NATIONALITY_TO_COUNTRY_OF_BIRTH[key]
    return ""


def _first_parsed_date(by_nk: Dict[str, Any], keys: Tuple[str, ...]) -> Optional[date]:
    for nk in keys:
        if nk in by_nk:
            d = _parse_flexible_date(by_nk[nk])
            if d is not None:
                return d
    return None


def _max_visa_stay_days_sane() -> int:
    """Reject absurd spans (wrong fields parsed as dates). Override via TURKEY_VISA_MAX_STAY_DAYS."""
    try:
        return max(31, min(3660, int(os.environ.get("TURKEY_VISA_MAX_STAY_DAYS", "370"))))
    except ValueError:
        return 370


def compute_inclusive_stay_days(flat: Dict[str, Any]) -> Optional[int]:
    """
    Inclusive calendar days for the trip.

    1) **Form-style:** `arrival_date` (enter) + `departure_date` or `return_date` (leave / fly back).
    2) **Zoho / sponsor-style (no arrival):** `departure_date` = trip start, `return_date` = trip end.

    Returns None if dates missing, unparseable, end before start, or span over sanity cap.
    """
    by_nk = _flat_first_by_normalized_key(flat)
    arr = _first_parsed_date(by_nk, _VISA_ARRIVAL_KEYS_ORDER)
    dep: Optional[date] = None

    if arr is None:
        d_go = _first_parsed_date(by_nk, ("departure_date", "departuredate"))
        d_back = _first_parsed_date(by_nk, ("return_date", "returndate", "date_of_return"))
        if d_go and d_back and d_back >= d_go:
            arr, dep = d_go, d_back
    else:
        dep = _first_parsed_date(by_nk, _VISA_END_AFTER_ARRIVAL_ORDER)

    if arr is None or dep is None:
        return None
    if dep < arr:
        return None
    days = (dep - arr).days + 1
    cap = _max_visa_stay_days_sane()
    if days > cap:
        return None
    return days


def build_replacements(flat: Dict[str, Any]) -> Dict[str, str]:
    cb = checkbox_placeholders(flat.get("sex"), flat.get("marital_status"))
    hist = turkey_history_checkboxes(flat)
    days = compute_inclusive_stay_days(flat)

    values: Dict[str, str] = {}
    for k, v in flat.items():
        nk = normalize_key(str(k)) if k is not None else ""
        if not nk:
            continue
        sval = _sanitize_for_word(str(v))
        values[nk] = sval

    for ck, cv in cb.items():
        values[ck] = cv
    for hk, hv in hist.items():
        values[hk] = hv

    # Always override body when dates yield a valid inclusive stay (Zoho often sends stale text).
    if days is not None:
        values["visa_duration"] = "1 day" if days == 1 else f"{days} days"

    # Derived from maid_nationality — do not use request `maid_country_of_birth` if sent.
    values["maid_country_of_birth"] = _sanitize_for_word(maid_country_of_birth_from_nationality(flat))

    # Word typo `{client)email}` → normalize_key gives `clientemail`, not `client_email`
    if "client_email" in values:
        values["clientemail"] = values["client_email"]

    # Templates use mixed case for R-Visa: `{RVisa_number}` and `{RVisa_expiry_date}`.
    # `normalize_key("RVisa_number")` → "r_visa_number" (split at the lowercase boundary),
    # but the request body uses `rvisa_number` → "rvisa_number". Bridge both forms.
    if "rvisa_number" in values:
        values["r_visa_number"] = values["rvisa_number"]
    if "rvisa_expiry_date" in values:
        values["r_visa_expiry_date"] = values["rvisa_expiry_date"]

    # Static rows: `{cbe}` / `{cbc}` (and aliases) — override JSON
    for rk, rv in _reserved_checkbox_glyph_placeholders().items():
        values[rk] = rv

    return values


def flatten_payload(body: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    fields = body.get("fields")
    if isinstance(fields, dict):
        out.update(fields)
    for k, v in body.items():
        if k == "fields" or v is None:
            continue
        out[k] = v
    return out


def _paragraph_base_run_font(paragraph):
    """
    Choose a font/size to apply to rebuilt runs in this paragraph.

    Strategy:
      1. If any run contains a `{` (placeholder start), use *that* run's font &
         size. This keeps substituted values (e.g. "62 days") at the same size
         as the placeholder was authored, so they don't suddenly become huge
         when the paragraph also contains a larger heading-style run such as
         the section number "23." or a checkbox glyph "\u2610".
      2. Otherwise fall back to the first run that has any explicit size.
      3. Otherwise fall back to the paragraph style.
    """
    placeholder_size = None
    placeholder_font = None
    first_size = None
    first_font = None

    for run in paragraph.runs:
        if not run.text:
            continue
        size = run.font.size
        name = run.font.name
        if "{" in run.text and placeholder_size is None and size is not None:
            placeholder_size = size
            placeholder_font = name
        if first_size is None and size is not None:
            first_size = size
        if not first_font and name:
            first_font = name

    chosen_size = placeholder_size or first_size
    chosen_font = placeholder_font or first_font

    if chosen_size is None:
        try:
            sf = paragraph.style.font.size
            if sf is not None:
                chosen_size = sf
        except (AttributeError, TypeError):
            pass
    if not chosen_font:
        try:
            fn = paragraph.style.font.name
            if fn:
                chosen_font = fn
        except (AttributeError, TypeError):
            pass
    return chosen_size, chosen_font


def _process_paragraph(paragraph, values: Dict[str, str]) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if "{" not in full_text:
        return

    base_size, base_font = _paragraph_base_run_font(paragraph)

    def repl(m: re.Match) -> str:
        key = normalize_key(m.group(1).strip())
        if key in values:
            return values[key]
        return m.group(0)

    segments: List[Tuple[str, bool]] = []
    pos = 0
    for m in _SINGLE_BRACE_RE.finditer(full_text):
        segments.append((full_text[pos : m.start()], False))
        key = normalize_key(m.group(1).strip())
        segments.append((repl(m), bool(key in values)))
        pos = m.end()
    segments.append((full_text[pos:], False))

    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    for text, _is_substituted in segments:
        if not text:
            continue
        r = paragraph.add_run(text)
        r.bold = is_bold

        # Don't force bold on substituted values: it makes filled fields stand
        # out from the rest of the form (e.g. "62 days" rendered larger/bolder
        # than its surrounding "Visa is requested for:" text). Inherit normal
        # weight from the paragraph style.
        if base_size is not None:
            r.font.size = base_size
        run_font = _font_name_for_run_text(text, base_font)
        if run_font:
            r.font.name = run_font


def _process_table(table, values: Dict[str, str]) -> None:
    """Walk a table and any nested tables (Word tables inside cells are not in doc.tables)."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _process_paragraph(p, values)
            for nested in cell.tables:
                _process_table(nested, values)


def _process_block(paragraphs, tables, values: Dict[str, str]) -> None:
    for p in paragraphs:
        _process_paragraph(p, values)
    for table in tables or []:
        _process_table(table, values)


def fill_turkey_docx_bytes(template_path: Path, flat: Dict[str, Any]) -> bytes:
    if not template_path.is_file():
        raise FileNotFoundError(f"Word template not found: {template_path}")

    values = build_replacements(flat)
    doc = Document(str(template_path))

    _process_block(doc.paragraphs, doc.tables, values)
    for section in doc.sections:
        for block in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            getattr(section, "even_page_header", None),
            getattr(section, "even_page_footer", None),
        ):
            if block is not None:
                _process_block(block.paragraphs, getattr(block, "tables", []) or [], values)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _resolve_soffice_path() -> Optional[str]:
    for name in ("soffice", "libreoffice"):
        p = shutil.which(name)
        if p:
            return p
    if os.name == "nt":
        for guess in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if Path(guess).is_file():
                return guess
    env = os.environ.get("LIBREOFFICE_PATH", "").strip()
    if env and Path(env).is_file():
        return env
    return None


def _convert_docx_bytes_to_pdf_soffice(docx_bytes: bytes) -> Optional[bytes]:
    soffice = _resolve_soffice_path()
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        td = Path(tmp)
        docx_path = td / "turkey_fill.docx"
        docx_path.write_bytes(docx_bytes)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(td), str(docx_path)],
                check=True,
                timeout=120,
                capture_output=True,
                text=True,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            return None
        pdf_path = td / "turkey_fill.pdf"
        if pdf_path.is_file():
            return pdf_path.read_bytes()
    return None


def _convert_docx_bytes_to_pdf_word_windows(docx_bytes: bytes) -> Optional[bytes]:
    """
    Convert using Microsoft Word via COM (Windows only). Requires Word installed.
    Optional: pip install pywin32
    """
    if os.name != "nt":
        return None
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        td = Path(tmp)
        docx_path = (td / "turkey_fill.docx").resolve()
        pdf_path = (td / "turkey_fill.pdf").resolve()
        docx_path.write_bytes(docx_bytes)
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                word.DisplayAlerts = 0
            except Exception:
                pass
            doc = word.Documents.Open(str(docx_path), ReadOnly=True)
            try:
                doc.SaveAs(str(pdf_path), FileFormat=17)
            finally:
                doc.Close(SaveChanges=False)
        except Exception:
            return None
        finally:
            if word is not None:
                try:
                    word.Quit(SaveChanges=False)
                except Exception:
                    pass
        if pdf_path.is_file():
            return pdf_path.read_bytes()
    return None


def convert_docx_bytes_to_pdf(docx_bytes: bytes) -> Optional[bytes]:
    """
    Convert filled .docx to .pdf: tries LibreOffice (soffice) first, then on Windows
    Microsoft Word (COM). Returns None if neither is available or conversion fails.
    """
    pdf = _convert_docx_bytes_to_pdf_soffice(docx_bytes)
    if pdf:
        return pdf
    return _convert_docx_bytes_to_pdf_word_windows(docx_bytes)


def list_single_brace_placeholders(doc_path: Path) -> List[str]:
    found: List[str] = []
    try:
        doc = Document(str(doc_path))
    except Exception:
        return found

    def scan_paragraphs(paragraphs) -> None:
        for paragraph in paragraphs:
            full_text = "".join(run.text for run in paragraph.runs)
            for m in _SINGLE_BRACE_RE.finditer(full_text):
                var_name = normalize_key(m.group(1).strip())
                if var_name and var_name not in found:
                    found.append(var_name)

    scan_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan_paragraphs(cell.paragraphs)
    for section in doc.sections:
        for block in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            getattr(section, "even_page_header", None),
            getattr(section, "even_page_footer", None),
        ):
            if block is not None:
                scan_paragraphs(block.paragraphs)
                for table in getattr(block, "tables", []) or []:
                    for row in table.rows:
                        for cell in row.cells:
                            scan_paragraphs(cell.paragraphs)

    return sorted(found)
