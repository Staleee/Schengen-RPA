"""
Documents Generation – replace {{variable}} placeholders in .docx templates.
Uses python-docx so the output is always a valid .docx that Word opens without errors.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document

# XML 1.0 invalid control chars (only \t \n \r are allowed in content)
_INVALID_XML_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\ufffe\uffff]")

# Placeholder in document text: {{variable_name}}
_PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")

# Substituted placeholder values are normally bold. A few placeholders carry
# ordinary sentence text (e.g. the NOC "or their family" clause) and must render
# in the same regular weight as the surrounding words, not bold.
NON_BOLD_PLACEHOLDERS = {"family_suffix"}

# Per-document variables (snake_case). Zoho will call each endpoint with its own body.
COVER_LETTER_VARIABLES = [
    "maid_full_name",
    "maid_passport_number",
    "schengen_country",  # Germany / France from Zoho
    "departure_date",
    "return_date",
    "client_name",
    "client_passport_number",
    "employment_start_date",  # contract start date from ERP
]

SPONSOR_LETTER_VARIABLES = [
    "client_name",
    "passport_number",
    "full_address_uae",
    "maid_full_name",
    "maid_passport_number",
    "employment_start_date",
    "salary_in_letters",
    "schengen_country",
    "departure_date",
    "return_date",
    "phone_number",  # client phone from Zoho
    "email",  # client email from Zoho
]

INVITATION_LETTER_VARIABLES = [
    "destination",  # Schengen country from Zoho
    "client_name",
    "address_in_uae",
    "maid_name",
    "contract_start_date",
    "arrival_date_to_departure_date",
    "cities",
    "hotel_address",
    "phone_number",
    "email_address",
]

DOCUMENT_VARIABLES = {
    "cover": COVER_LETTER_VARIABLES,
    "sponsor": SPONSOR_LETTER_VARIABLES,
    "invitation": INVITATION_LETTER_VARIABLES,
}


def normalize_key(text: str) -> str:
    """Normalize to snake_case: camelCase, strip, lowercase, spaces -> underscores."""
    if not text or not isinstance(text, str):
        return ""
    t = text.strip()
    # camelCase / PascalCase -> snake_case (e.g. departureDate, Return_Date -> departure_date)
    t = re.sub(r"([a-z0-9])([A-Z])", r"\1_\2", t)
    t = re.sub(r"([A-Z]+)([A-Z][a-z])", r"\1_\2", t)
    t = t.lower()
    t = re.sub(r"\s+", "_", t)
    t = re.sub(r"[^\w\-]", "", t)
    return t


def _sanitize_for_word(value: str) -> str:
    """Strip characters that break Word (control chars, invalid Unicode)."""
    if not value:
        return ""
    return _INVALID_XML_RE.sub("", value)


_SCHENGEN_TURKEY_RE = re.compile(r"Schengen\s+Turkey|Turkey\s+Schengen", re.IGNORECASE)
_SCHENGEN_WORD_RE = re.compile(r"\bSchengen\b", re.IGNORECASE)
# After substitution + Schengen→Turkey, "{{schengen_country}} Schengen" becomes
# "Turkey Turkey" when Turkey is the country. Collapse those duplicates.
_DUPLICATE_TURKEY_RE = re.compile(r"\bTurkey(?:\s+Turkey)+\b", re.IGNORECASE)

# Spain submits Schengen visas through BLS instead of VFS. The sponsor letter
# template hard-codes "VFS"; swap it for "BLS" when the destination is Spain.
# Case-sensitive on purpose to avoid touching unrelated lowercase tokens.
_VFS_WORD_RE = re.compile(r"\bVFS\b")


def _replace_schengen_with_turkey(text: str) -> str:
    """Remove combined 'Schengen Turkey'/'Turkey Schengen' phrases, then replace remaining 'Schengen' with 'Turkey', then collapse any 'Turkey Turkey' duplicates."""
    text = _SCHENGEN_TURKEY_RE.sub("Turkey", text)
    text = _SCHENGEN_WORD_RE.sub("Turkey", text)
    text = _DUPLICATE_TURKEY_RE.sub("Turkey", text)
    return text


def _replace_vfs_with_bls(text: str) -> str:
    """Replace standalone 'VFS' tokens with 'BLS' (Spain submission channel)."""
    return _VFS_WORD_RE.sub("BLS", text)


# Match the consulate/embassy phrase + an optional "– <city/country>" tail.
# Examples that should match (after placeholder substitution):
#   "Embassy of Turkey"
#   "Embassy of Turkey – UAE"
#   "Consulate of Turkey – Dubai"
#   "Embassy/Consulate of Turkey – UAE"
# The optional tail is greedy enough to swallow "– UAE", "– Abu Dhabi", etc.
_EMBASSY_PHRASE_RE = re.compile(
    r"(?:Embassy\s*/\s*Consulate|Embassy|Consulate)\s+of\s+[A-Za-z][A-Za-z\-]*"
    r"(?:\s*[\u2013\-]\s*[A-Za-z][A-Za-z ,]*)?",
    re.IGNORECASE,
)

# Stand-alone city paragraph (the templates carry "Abu Dhabi, United Arab
# Emirates" on its own line below the embassy line). When the addressee block
# is replaced we either swap this line for the new location or clear it,
# depending on whether the caller sent a two-line or legacy single-line
# addressee.
_CITY_LINE_RE = re.compile(r"^\s*(Abu\s*Dhabi|Dubai)(\s*[,\-\u2013].*)?$", re.IGNORECASE)


def _replace_embassy_block_with_addressee(
    paragraphs_in_order,
    addressee: str,
    location: str = "",
) -> None:
    """Find the first paragraph containing an "Embassy of <country>" / "Consulate
    of <country>" / "Embassy/Consulate of <country>" phrase (including any
    trailing "– <city>") and substitute that phrase with ``addressee``.

    Also handles the standalone city line that lives on its own paragraph below
    (e.g. "Abu Dhabi, United Arab Emirates"):

      * When ``location`` is provided, the city paragraph is replaced with the
        new location text. This is the two-line addressee mode used for all
        Schengen countries — the addressee body goes on line 1 and the city /
        country line goes on line 2 ("Consulate of Croatia" / "Dubai, United
        Arab Emirates").
      * When ``location`` is empty (legacy single-line Turkey mode), the city
        paragraph is cleared instead so the document doesn't show a stale city.
    """
    if not addressee:
        return
    for i, para in enumerate(paragraphs_in_order):
        full_text = "".join(run.text for run in para.runs)
        match = _EMBASSY_PHRASE_RE.search(full_text)
        if not match:
            continue

        new_text = full_text[: match.start()] + addressee + full_text[match.end() :]
        for run in list(para.runs):
            run._r.getparent().remove(run._r)
        para.add_run(new_text)

        # Update / clear the immediately-following standalone city paragraph
        # only when the original embassy paragraph was a clean line (no other
        # text around the matched phrase). Otherwise the city info likely lives
        # on a later line of the same paragraph and our replacement already
        # handled it.
        if match.start() == 0 and match.end() == len(full_text.rstrip()):
            for offset in (1, 2):
                j = i + offset
                if j >= len(paragraphs_in_order):
                    break
                next_para = paragraphs_in_order[j]
                next_text = "".join(run.text for run in next_para.runs).strip()
                if not next_text:
                    continue
                if _CITY_LINE_RE.match(next_text):
                    for run in list(next_para.runs):
                        run._r.getparent().remove(run._r)
                    if location:
                        next_para.add_run(location)
                break
        return  # Only replace the first match.


def fill_document(doc_path: Path, variables: Dict[str, str], output_path: Path) -> List[str]:
    """
    Replace {{variable_name}} placeholders using python-docx. Output is always a valid
    .docx that Word opens without "unspecified error". Handles placeholders split across runs.
    When variables contain is_turkey=True/true/1, all occurrences of "Schengen" (and the
    combined phrase "Schengen Turkey"/"Turkey Schengen") in the document body are replaced
    with "Turkey" after placeholder substitution.
    When variables contain is_spain=True/true/1, the sponsor letter's hard-coded "VFS" is
    rewritten to "BLS" (Spain Schengen visas are submitted through BLS, not VFS).
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not variables:
        import shutil
        shutil.copy2(doc_path, output_path)
        return []

    normalized = {normalize_key(k): _sanitize_for_word((str(v) if v is not None else "").strip()) for k, v in variables.items()}
    filled: List[str] = []

    # Determine Turkey / Spain mode from the mapped variables.
    _is_turkey = normalized.get("is_turkey", "").lower() in ("true", "1", "yes")
    _is_spain = normalized.get("is_spain", "").lower() in ("true", "1", "yes")

    def repl(match: re.Match) -> str:
        var_name = normalize_key(match.group(1).strip())
        if var_name == "date":
            # {{date}} = current date when the document is generated (not from request body)
            filled.append("date")
            n = datetime.now()
            return f"{n.month}/{n.day}/{n.year}"  # e.g. 3/10/2026
        if var_name in normalized:
            filled.append(var_name)
            return normalized[var_name]
        return match.group(0)

    def process_paragraph(paragraph):
        """Replace placeholders; only the substituted variable values are bold."""
        full_text = "".join(run.text for run in paragraph.runs)
        has_placeholder = "{{" in full_text
        has_schengen = _is_turkey and ("Schengen" in full_text or "schengen" in full_text)
        has_vfs = _is_spain and "VFS" in full_text

        if not has_placeholder and not has_schengen and not has_vfs:
            return

        if has_placeholder:
            # Build segments: (text, bold) – bold only for substituted values
            segments = []
            pos = 0
            for m in _PLACEHOLDER_RE.finditer(full_text):
                segments.append((full_text[pos : m.start()], False))
                var_name = normalize_key(m.group(1).strip())
                segments.append((repl(m), var_name not in NON_BOLD_PLACEHOLDERS))
                pos = m.end()
            segments.append((full_text[pos:], False))
        else:
            segments = [(full_text, False)]

        if _is_turkey:
            # Apply Schengen→Turkey on the JOINED text so cross-segment patterns like
            # "{{schengen_country}}" + " Schengen Visa" → "Turkey" + " Turkey Visa" get
            # collapsed to "Turkey Visa". Bold formatting is preserved when the cleaned
            # text is identical to the joined original.
            joined = "".join(t for t, _ in segments)
            cleaned = _replace_schengen_with_turkey(joined)
            if cleaned != joined:
                segments = [(cleaned, False)]

        if _is_spain:
            # Spain uses BLS for Schengen visa submission. The hard-coded "VFS"
            # in the sponsor letter sits inside a non-bold run, so we only rewrite
            # the per-segment text and keep the original bold flags intact.
            new_segments = []
            changed = False
            for text, is_bold in segments:
                rewritten = _replace_vfs_with_bls(text)
                if rewritten != text:
                    changed = True
                new_segments.append((rewritten, is_bold))
            if changed:
                segments = new_segments

        # Clear and rebuild: one run per segment, bold only for substituted values
        for run in list(paragraph.runs):
            run._r.getparent().remove(run._r)
        for text, is_bold in segments:
            if not text:
                continue
            r = paragraph.add_run(text)
            r.bold = is_bold

    def process_block(paragraphs, tables=None):
        for p in paragraphs:
            process_paragraph(p)
        if tables:
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            process_paragraph(p)

    doc = Document(doc_path)

    # Body: paragraphs and tables
    process_block(doc.paragraphs, doc.tables)

    # Replace the "Embassy of <country>" + city block in the To: section so the
    # letter addresses the office that will actually process the visa
    # application (Embassy in Abu Dhabi vs. Consulate in Dubai). This applies to
    # all Schengen countries — the caller decides Embassy vs. Consulate from the
    # client's UAE address and passes the result via:
    #
    #   * addressee_body     — line 1, e.g. "Embassy of Croatia" / "Consulate of Croatia"
    #   * addressee_location — line 2, e.g. "Abu Dhabi, United Arab Emirates" / "Dubai, United Arab Emirates"
    #
    # The legacy single-line ``addressee`` (e.g. "Consulate of Turkey – Dubai")
    # is still accepted for backwards compatibility with older callers; in that
    # mode the city line below is cleared rather than replaced.
    addressee_body = normalized.get("addressee_body", "").strip()
    addressee_location = normalized.get("addressee_location", "").strip()
    legacy_addressee = normalized.get("addressee", "").strip()

    if addressee_body or legacy_addressee:
        ordered_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    ordered_paragraphs.extend(cell.paragraphs)
        if addressee_body:
            _replace_embassy_block_with_addressee(
                ordered_paragraphs, addressee_body, addressee_location
            )
        else:
            _replace_embassy_block_with_addressee(ordered_paragraphs, legacy_addressee)

    # Headers and footers (where maid_full_name at bottom/signature often lives)
    for section in doc.sections:
        for block in (section.header, section.footer, section.first_page_header, section.first_page_footer, getattr(section, "even_page_header", None), getattr(section, "even_page_footer", None)):
            if block is not None:
                process_block(block.paragraphs, getattr(block, "tables", []))

    doc.save(str(output_path))
    return filled


def list_placeholder_variables(doc_path: Path) -> List[str]:
    """List {{variable}} placeholders in the document (paragraphs and table cells)."""
    found: List[str] = []
    try:
        doc = Document(doc_path)
    except Exception:
        return found
    for paragraph in doc.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        for m in _PLACEHOLDER_RE.finditer(full_text):
            var_name = normalize_key(m.group(1).strip())
            if var_name and var_name not in found:
                found.append(var_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = "".join(run.text for run in paragraph.runs)
                    for m in _PLACEHOLDER_RE.finditer(full_text):
                        var_name = normalize_key(m.group(1).strip())
                        if var_name and var_name not in found:
                            found.append(var_name)
    return sorted(found)
