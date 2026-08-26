"""One-time import: install the business-provided Turkey client NOC Word template.

Source: ``turkey/NOC (1).docx`` (provided by operations). It uses the old @token@ style; this
converts the tokens to the house {{placeholder}} style with the standard request-key names and
writes ``client-noc.docx`` next to the other letter templates.

Token mapping:
  @client_name@               → {{client_name}}
  @client_nationality@        → {{client_nationality}}
  @Client_passport_number@    → {{client_passport_number}}
  @Housemaid_name@            → {{maid_name}}
  @Housemaid_nationality@     → {{maid_nationality}}
  @Housemaid_passport_number@ → {{maid_passport_number}}
  @Client_phone_number@       → {{client_contact_number}}

The conversion is character-by-character so per-run formatting survives. An earlier version
rebuilt the whole paragraph into its first run whenever a token straddled a run boundary — which
every token in this document does, because Google Docs splits runs mid-word. That flattened the
one run that carries meaning: "Wife / Husband /Son / Daughter /" is struck through so that
"Maid" reads as the operative word, and the generated letter lost the strike entirely.

Idempotent: re-running regenerates client-noc.docx from the same source.

    python scripts/import_client_noc_docx.py
"""

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent.parent
SOURCE = BASE_DIR / "turkey" / "NOC (1).docx"
TARGET = BASE_DIR / "client-noc.docx"

# Tokens land on the REQUEST key names — the docx fill matches placeholders by request key.
TOKEN_MAP = {
    "client_name": "client_name",
    "client_nationality": "client_nationality",
    "Client_passport_number": "client_passport_number",
    "Housemaid_name": "maid_name",
    "Housemaid_nationality": "maid_nationality",
    "Housemaid_passport_number": "maid_passport_number",
    "Client_phone_number": "client_contact_number",
}

_TOKEN_RE = re.compile("@(" + "|".join(re.escape(t) for t in TOKEN_MAP) + ")@")


def _run_props(run):
    """The run's properties, minus a stray underline the source carries.

    The provided document has exactly one underlined run: the single letter "h" of "holding",
    left behind by an editing slip in Google Docs. It used to vanish because the import flattened
    all formatting; now that formatting is preserved, it would print as a stray mark under one
    letter. No underline in this letter is intentional, so they are all dropped.
    """
    r_pr = run._r.find(qn("w:rPr"))
    if r_pr is None:
        return None
    for underline in r_pr.findall(qn("w:u")):
        r_pr.remove(underline)
    return r_pr


def _center_heading(para) -> bool:
    """Centre a heading Google Docs exported as right-aligned-plus-right-indent.

    The two headings are meant to sit centred. Google Docs expresses that as ``jc="right"`` with
    a large right indent, which is not centring at all: it narrows the line's usable width, and
    "To The Turkish Embassy / Consulate, UAE" no longer fitted, so "UAE" wrapped onto a line of
    its own. Real centring with the indent cleared gives the full column width and the intended
    look.
    """
    p_pr = para._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    jc = p_pr.find(qn("w:jc"))
    if jc is None or jc.get(qn("w:val")) != "right":
        return False
    jc.set(qn("w:val"), "center")
    ind = p_pr.find(qn("w:ind"))
    if ind is not None:
        p_pr.remove(ind)
    return True


def _convert_paragraph(para) -> int:
    """Rewrite @tokens@ as {{placeholders}}, keeping every run's own formatting.

    Works on a character-to-run map rather than on run text, because a token routinely spans
    runs. Static text keeps the formatting of the characters it came from; a placeholder takes
    the formatting of its token's last character, which is the run the author styled the token
    with (the leading "@" often belongs to the preceding plain run).
    """
    runs = list(para.runs)
    if not runs:
        return 0
    text = "".join(run.text for run in runs)
    if not _TOKEN_RE.search(text):
        return 0

    props = [_run_props(run) for run in runs]
    origin = []
    for index, run in enumerate(runs):
        origin.extend([props[index]] * len(run.text))

    # (text, run properties) in order, tokens replaced.
    pieces = []
    pos = 0
    for match in _TOKEN_RE.finditer(text):
        start, end = match.span()
        i = pos
        while i < start:  # static text, split on formatting changes
            fmt = origin[i]
            j = i
            while j < start and origin[j] is fmt:
                j += 1
            pieces.append((text[i:j], fmt))
            i = j
        pieces.append(("{{" + TOKEN_MAP[match.group(1)] + "}}", origin[end - 1]))
        pos = end
    i = pos
    while i < len(text):
        fmt = origin[i]
        j = i
        while j < len(text) and origin[j] is fmt:
            j += 1
        pieces.append((text[i:j], fmt))
        i = j

    for run in runs:
        run._r.getparent().remove(run._r)
    for piece_text, fmt in pieces:
        if not piece_text:
            continue
        new_run = para.add_run(piece_text)
        if fmt is not None:
            new_run._r.insert(0, copy.deepcopy(fmt))
    return 1


def main() -> None:
    if not SOURCE.exists():
        print(f"  ERROR: source not found: {SOURCE}")
        return

    doc = Document(str(SOURCE))
    converted = 0
    centred = 0
    for para in doc.paragraphs:
        converted += _convert_paragraph(para)
        centred += 1 if _center_heading(para) else 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    converted += _convert_paragraph(para)

    doc.save(str(TARGET))
    print(f"  converted {converted} paragraph(s), centred {centred} heading(s); wrote {TARGET.name}")


if __name__ == "__main__":
    main()
