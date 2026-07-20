"""
One-time script: two edits to noc-travel.docx.

  1. Replace the sentence-ending period after the companion passport placeholder
     with a {{family_suffix}} placeholder:
        ...holding passport number {{companion_passport_number}}.
     becomes
        ...holding passport number {{companion_passport_number}}{{family_suffix}}
     The Next.js payload fills family_suffix with ", or their family." for Egypt
     NOCs and just "." for every other flow (Lebanon / Schengen). The engine
     renders the family_suffix value non-bold (see NON_BOLD_PLACEHOLDERS in
     doc_utils.py) so it matches the surrounding sentence.

  2. Insert a new paragraph immediately after the "Visa Type: {{visa_type}}" line
     holding the placeholder {{stay_line}}. The payload fills this with
     " - Stay: 2 months" only for Egypt + Single-entry + Egyptian-companion NOCs
     where the maid requested the one-month extension; otherwise it is sent empty
     and the line renders blank.

The fill engine (doc_utils.fill_document) does plain {{placeholder}} substitution
with no conditionals, which is why both country-specific pieces are driven by
placeholder values rather than template logic.

Idempotent: safe to run more than once (skips edits already applied).

Run once from the Documents Generation directory:
    python scripts/patch_noc_family_and_stay.py
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE_DIR / "noc-travel.docx"

PASSPORT_PLACEHOLDER = "{{companion_passport_number}}"
FAMILY_PLACEHOLDER = "{{family_suffix}}"
STAY_PLACEHOLDER = "{{stay_line}}"


def _set_paragraph_text(para, new_text: str) -> None:
    """Assign the whole text to the first run and clear the rest (keeps run 0
    formatting; matches the technique used by patch_noc_placeholders.py)."""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _patch_family_suffix(doc) -> bool:
    """Swap the '.' after {{companion_passport_number}} for {{family_suffix}}."""
    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs)
        if PASSPORT_PLACEHOLDER not in full_text:
            continue
        if FAMILY_PLACEHOLDER in full_text:
            print("  family_suffix placeholder already present - skipping")
            return False
        if f"{PASSPORT_PLACEHOLDER}." in full_text:
            new_text = full_text.replace(
                f"{PASSPORT_PLACEHOLDER}.",
                f"{PASSPORT_PLACEHOLDER}{FAMILY_PLACEHOLDER}",
                1,
            )
        else:
            # No trailing period; just append the placeholder after the passport.
            new_text = full_text.replace(
                PASSPORT_PLACEHOLDER,
                f"{PASSPORT_PLACEHOLDER}{FAMILY_PLACEHOLDER}",
                1,
            )
        _set_paragraph_text(para, new_text)
        print("  family_suffix placeholder added")
        return True
    print("  WARNING: companion passport placeholder not found")
    return False


def _patch_stay_line(doc) -> bool:
    """Insert a '{{stay_line}}' paragraph right after the Visa Type paragraph."""
    for para in doc.paragraphs:
        if STAY_PLACEHOLDER in "".join(run.text for run in para.runs):
            print("  stay line placeholder already present - skipping")
            return False

    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs)
        if "Visa Type:" not in full_text:
            continue
        # Clone the Visa Type paragraph element to inherit its formatting (pPr),
        # then strip its runs and set the stay placeholder as the only run.
        new_p = copy.deepcopy(para._p)
        for r in new_p.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
        ):
            new_p.remove(r)
        para._p.addnext(new_p)
        from docx.text.paragraph import Paragraph

        new_para = Paragraph(new_p, para._parent)
        new_para.add_run(STAY_PLACEHOLDER)
        print("  stay line placeholder inserted after Visa Type")
        return True
    print("  WARNING: 'Visa Type:' paragraph not found")
    return False


def patch_template(path: Path) -> None:
    if not path.exists():
        print(f"  ERROR: template not found at {path}")
        return

    doc = Document(str(path))
    changed = False
    changed |= _patch_family_suffix(doc)
    changed |= _patch_stay_line(doc)

    if changed:
        doc.save(str(path))
        print(f"  SAVED {path.name}")
    else:
        print(f"  OK (nothing to change): {path.name}")


if __name__ == "__main__":
    print("Patching NOC travel template (family_suffix + stay line)...")
    patch_template(TEMPLATE_PATH)
    print("Done.")
