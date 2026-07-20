# -*- coding: utf-8 -*-
"""
One-time script: reword the three visa letters so the domestic worker is no
longer described as employed under the *client's* personal sponsorship. The
employment is attributed to Maids.cc instead.

Replacements (each target phrase currently sits inside a single run in every
template, so we replace at run level to keep the {{placeholders}} and bold
formatting intact):

  Sponsor_letter.docx
    "under my sponsorship"          -> "under Maids.cc"
    "employed with me"              -> "working at my house"

  Cover_Letter.docx
    "under the sponsorship of"      -> "under Maids.cc working at the house of"
      (the following {{client_name}} placeholder is left untouched, so it reads
       "under Maids.cc working at the house of {{client name}}")

  Invitation Letter (Schengen Visa - Domestic Worker_Housemaid).docx
    "employed under my sponsorship" -> "employed under Maids.cc and working at my house"

Idempotent: replacements only fire when the old phrase is still present, so
re-running is a no-op. A warning is printed if a target phrase is not found
(e.g. a template was re-saved and split the phrase across runs).

Run once from the Documents Generation directory:
    python scripts/patch_letters_maidscc_sponsor.py
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent

# filename -> list of (old_phrase, new_phrase)
REPLACEMENTS: dict[str, list[tuple[str, str]]] = {
    "Sponsor_letter.docx": [
        ("under my sponsorship", "under Maids.cc"),
        ("employed with me", "working at my house"),
    ],
    "Cover_Letter.docx": [
        ("under the sponsorship of", "under Maids.cc working at the house of"),
    ],
    "Invitation Letter (Schengen Visa – Domestic Worker_Housemaid).docx": [
        (
            "employed under my sponsorship",
            "employed under Maids.cc and working at my house",
        ),
    ],
}


def _iter_all_paragraphs(doc) -> Iterable:
    """Yield every paragraph in the body, tables and headers/footers."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    for section in doc.sections:
        for blk in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            getattr(section, "even_page_header", None),
            getattr(section, "even_page_footer", None),
        ):
            if blk is not None:
                for para in blk.paragraphs:
                    yield para


def patch_file(path: Path, replacements: list[tuple[str, str]]) -> None:
    if not path.exists():
        print(f"  ERROR: template not found: {path.name}")
        return

    doc = Document(str(path))
    counts = {old: 0 for old, _ in replacements}

    for para in _iter_all_paragraphs(doc):
        for run in para.runs:
            for old, new in replacements:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    counts[old] += 1

    applied = sum(counts.values())
    if applied:
        doc.save(str(path))
    for old, new in replacements:
        n = counts[old]
        if n:
            print(f"  [{path.name}] {n}x  {old!r} -> {new!r}")
        else:
            # Already done, or the phrase moved across runs.
            print(f"  [{path.name}] SKIP (not found): {old!r}")


if __name__ == "__main__":
    print("Rewording visa letters: client sponsorship -> Maids.cc ...")
    for filename, reps in REPLACEMENTS.items():
        patch_file(BASE_DIR / filename, reps)
    print("Done.")
