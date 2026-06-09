"""
One-time script: converts @placeholder@ tokens in noc-travel.docx to {{placeholder}}
so the existing fill_document engine can process them.

Mapping applied:
  @today_date@                → {{date}}          (auto-filled by fill_document)
  @maid_name@                 → {{maid_name}}
  @maid_passport_number@      → {{maid_passport_number}}
  @maid_nationality@          → {{maid_nationality}}
  @maid_eid_number@           → {{maid_eid_number}}
  @maid_profession@           → {{maid_profession}}
  @destination_country@       → {{destination_country}}
  @companion_salutation_name@ → {{companion_salutation_name}}
  @companion_name@            → {{companion_name}}
  @companion_nationality@     → {{companion_nationality}}
  @companion_passport_number@ → {{companion_passport_number}}
  @visa_type@                 → {{visa_type}}

Run once from the Documents Generation directory:
    python scripts/patch_noc_placeholders.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE_DIR / "noc-travel.docx"

# Explicit mapping for tokens that differ from the default @x@ → {{x}} rule.
_EXPLICIT: dict[str, str] = {
    "today_date": "date",
}

_AT_RE = re.compile(r"@([^@]+)@")


def _convert(token: str) -> str:
    """Return the {{placeholder}} string for a given @token@ inner value."""
    return "{{" + _EXPLICIT.get(token, token) + "}}"


def _patch_text(text: str) -> str:
    return _AT_RE.sub(lambda m: _convert(m.group(1)), text)


def _patch_paragraph(para) -> bool:
    """Rebuild runs in para so that @token@ → {{placeholder}}. Returns True if changed."""
    full_text = "".join(run.text for run in para.runs)
    new_text = _patch_text(full_text)
    if new_text == full_text:
        return False

    # Assign entire new text to the first run, clear the rest.
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    return True


def patch_template(path: Path) -> None:
    if not path.exists():
        print(f"  ERROR: template not found at {path}")
        return

    doc = Document(str(path))
    changed = 0

    for para in doc.paragraphs:
        if _patch_paragraph(para):
            changed += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _patch_paragraph(para):
                        changed += 1

    # Headers / footers
    for section in doc.sections:
        for hdr_ftr in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hdr_ftr is not None:
                for para in hdr_ftr.paragraphs:
                    if _patch_paragraph(para):
                        changed += 1

    if changed:
        doc.save(str(path))
        print(f"  PATCHED {changed} paragraph(s) in {path.name}")
    else:
        print(f"  OK (no @placeholder@ tokens found): {path.name}")


if __name__ == "__main__":
    print("Patching NOC travel template placeholders...")
    patch_template(TEMPLATE_PATH)
    print("Done.")
