"""
One-time script: insert a {{addressee}} placeholder paragraph at the top of
noc-travel.docx so the certificate can be addressed to a specific consulate
(e.g. the Saudi Consulate) when the caller provides one.

The fill engine (doc_utils.fill_document) does plain {{placeholder}}
substitution with no conditionals, so the addressee is driven entirely by the
payload value: the backend sends the addressee text only for countries that
require one (Saudi Arabia) and sends an empty string otherwise, in which case
the line renders blank.

Idempotent: safe to run more than once (skips when already present).

Run once from the Documents Generation directory:
    python scripts/patch_noc_addressee.py
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE_DIR / "noc-travel.docx"

ADDRESSEE_PLACEHOLDER = "{{addressee}}"


def _patch_addressee(doc) -> bool:
    """Insert a '{{addressee}}' paragraph immediately before the Subject line."""
    for para in doc.paragraphs:
        if ADDRESSEE_PLACEHOLDER in "".join(run.text for run in para.runs):
            print("  addressee placeholder already present - skipping")
            return False

    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs)
        if "Subject:" not in full_text:
            continue
        # Clone the Subject paragraph element to inherit its formatting (pPr),
        # strip its runs, set the addressee placeholder as the only run, and
        # insert it before the Subject paragraph.
        new_p = copy.deepcopy(para._p)
        for r in new_p.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
        ):
            new_p.remove(r)
        para._p.addprevious(new_p)
        new_para = Paragraph(new_p, para._parent)
        new_para.add_run(ADDRESSEE_PLACEHOLDER)
        print("  addressee placeholder inserted before Subject")
        return True
    print("  WARNING: 'Subject:' paragraph not found")
    return False


def patch_template(path: Path) -> None:
    if not path.exists():
        print(f"  ERROR: template not found at {path}")
        return

    doc = Document(str(path))
    if _patch_addressee(doc):
        doc.save(str(path))
        print(f"  SAVED {path.name}")
    else:
        print(f"  OK (nothing to change): {path.name}")


if __name__ == "__main__":
    print("Patching NOC travel template (addressee)...")
    patch_template(TEMPLATE_PATH)
    print("Done.")
