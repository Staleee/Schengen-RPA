"""Build ``noc-schengen.docx`` — the two-page (English + Arabic) maid NOC letter.

Replaces ``Travel_NOC_Fillable.pdf`` as the template behind document_type=noc-schengen
and noc-turkey. The AcroForm could not work: it is a flat letter with fixed-width form
fields laid over the gaps, so a value cannot reflow the sentence around it. Real data
overflowed those gaps — a 30-character maid name got shrunk to 5pt and clipped mid-word
("Kusnayati BT Sukri"). A .docx goes through the same fill + LibreOffice path as every
other letter here, where text reflows and Word shapes the Arabic itself.

That PDF has since been removed from the tree (recover it from git history if the wording ever
needs re-checking against the original). Its content was recovered rather than retyped:

* The letterhead PNG was extracted from it and is committed at assets/maidscc_letterhead.png.
* The English page is taken verbatim from the PDF's static text.
* The Arabic page was de-shaped back to base letters and is embedded below as literals.
  Every Arabic glyph in the PDF is a presentation form (U+FB50..U+FEFF) that NFKC maps
  losslessly to its base letter, already in logical order, so the recovery was exact:
  ``"".join(unicodedata.normalize("NFKC", c) for c in span_text)``. Punctuation the PDF
  stored at the visual edge of a run was moved back to its logical position by hand.

Three defects in the PDF template are repaired on the way, all verified present in it:

1. The Arabic "Emirates ID" phrase was shattered across two lines as
   "Emirates" / ". المتحدة" / "م" — rewritten as الهوية الإماراتية.
2. "Aldrin Dumigpi, Employee Relations Officer" was drawn at the exact same position as
   "HR Manager" on the Arabic page — two signature blocks stacked on each other. Dropped.
3. The English trip-expenses line read "·s expenses" (a mangled apostrophe) — now "'s".

The Arabic wording is a reconstruction and needs a native-speaker review before this ships.

    python scripts/build_noc_schengen_docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT = BASE_DIR / "noc-schengen.docx"
LOGO = BASE_DIR / "assets" / "maidscc_letterhead.png"

# The letterhead sits at [181, 36, 431, 126] on both PDF pages — 250pt wide.
LOGO_WIDTH_PT = 250

BODY_FONT = "Times New Roman"
# Complex-script (Arabic) face, named explicitly rather than left to font fallback. The Latin
# face is Times New Roman, which LibreOffice maps to Liberation Serif — a family with no Arabic
# glyphs at all, so relying on the Latin name left the Arabic page as empty boxes. DejaVu Sans
# is the Arabic-capable font present in the service image (see Dockerfile).
ARABIC_FONT = "DejaVu Sans"
BODY_SIZE = 11
TITLE_FONT = "Arial"
TITLE_SIZE = 14

FOOTER_ATTESTATION = (
    "This is an original attested digital certificate. "
    "No other form of this certificate can be generated."
)
FOOTER_ADDRESS_AR = (
    "ميدز دوت سي سي، صندوق بريد P.O. Box 40398، ملك محمد صالح عبدالله ال علي، "
    "معرض رقم 1، البرشاء الثانية، الإمارات العربية المتحدة | رقم الاتصال:"
)
FOOTER_PHONE = "04-2479155"
FOOTER_DISCLAIMER = (
    "This No-Objection Certificate (NOC) is issued solely for the purpose specified herein "
    "and does not imply any legal or financial liability on the part of Maids.cc for any "
    "actions, conduct, or obligations of the individual beyond the scope of their designated "
    "role as defined by their contractual obligations with Maids.cc. Maids.cc expressly "
    "disclaims all liability for any misuse, unauthorized use, or misrepresentation of this "
    "certificate. Any such actions are the sole responsibility of the individual and/or their "
    "sponsor, regardless of the purpose for which the NOC is issued"
)

ENGLISH_TITLE = "Non-Objection Certificate"
ARABIC_TITLE = "شهادة لا مانع سفر"

# (style, text) — "h" = bold heading, "b" = body, "l" = bullet/dash item.
ENGLISH_BODY = [
    ("b", "{{date_issued}}"),
    ("b", "To the valued staff of the {{consulate_full_name}},"),
    ("s", ""),
    (
        "b",
        "This is to confirm that Maids.cc has no objection for {{worker_name}} to travel "
        "to {{destination_country}} with the mentioned traveller(s) below:",
    ),
    ("s", ""),
    ("h", "Traveller(s) Details:"),
    ("b", "{{worker_name}} will be travelling with,"),
    # The companion's passport / EID are optional on the application, so the clause is
    # composed in variable_enrichment rather than printing labels around empty values.
    ("l", "{{companion_name}} {{companion_id_clause}}"),
    ("s", ""),
    ("h", "Domestic Worker Details:"),
    (
        "b",
        "{{worker_name}}, of the {{worker_nationality}} nationality and holder of passport "
        "number {{worker_passport}} and Emirates ID: {{worker_eid}}, has been employed by "
        "Maids.cc since {{employment_start_date}} until the present day, and has an active "
        "visa that expires on {{visa_expiry_date}}. {{worker_name}} is assigned to work in "
        "the household of {{employer_name}} on a {{employment_basis}} basis. {{worker_name}} "
        "receives a monthly total salary of AED {{monthly_salary}} per month, equivalent to "
        "AED {{annual_salary}} per year.",
    ),
    ("s", ""),
    ("h", "Trip Details:"),
    ("l", "Trip Purpose: {{trip_purpose}}"),
    ("l", "Proposed Travel Dates: {{travel_start_date}} to {{travel_end_date}}"),
    (
        "l",
        "Trip Expenses: {{companion_name}} will pay for {{worker_name}}'s expenses "
        "throughout her trip to {{destination_country}}.",
    ),
    ("s", ""),
    ("h", "Respectfully submitted,"),
    ("b", "HR Manager"),
    ("b", "+971 505544143"),
]

# Same letter, Arabic page. employment_basis / trip_purpose arrive in English, so the
# Arabic page states them as static Arabic text (أساس طويل الأمد / السياحة) exactly as the
# PDF template did — the placeholders stay on the English page only.
ARABIC_BODY = [
    ("b", "{{date_issued}}"),
    ("b", "إلى الموظفين الكرام في سفارة {{destination_country}} بالإمارات العربية المتحدة،"),
    ("s", ""),
    (
        "b",
        "تؤكد شركة ميدز دوت سي سي لخدمات العمالة المساعدة ش.ذ.م.م بأنه ليس لدينا مانع من سفر "
        "{{worker_name}} إلى {{destination_country}} مع المرافقين المذكورين أدناه.",
    ),
    ("s", ""),
    ("h", "معلومات المرافقين المسافر:"),
    ("b", "ستسافر {{worker_name}} مع:"),
    ("l", "{{companion_name}} {{companion_id_clause_ar}}"),
    ("s", ""),
    ("h", "معلومات الخادمة:"),
    (
        "b",
        "تم توظيف {{worker_name}} من الجنسية {{worker_nationality}} وحاملة رقم جواز السفر "
        "{{worker_passport}} والهوية الإماراتية {{worker_eid}}، لدى شركة ميدز دوت سي سي "
        "لخدمات العمالة المساعدة ش.ذ.م.م منذ {{employment_start_date}} حتى يومنا هذا وهي "
        "بموجب تأشيرة سارية تنتهي في {{visa_expiry_date}}. تم تعيين {{worker_name}} للعمل "
        "في منزل {{employer_name}} على أساس طويل الأمد. تتقاضى {{worker_name}} راتبًا "
        "اجماليًا شهريًا قدره {{monthly_salary}} درهم شهريًا، أي ما يعادل {{annual_salary}} "
        "درهم في السنة.",
    ),
    ("s", ""),
    ("h", "معلومات الرحلة:"),
    ("l", "الغرض من الرحلة: السياحة"),
    ("l", "مواعيد السفر المتاحة: {{travel_start_date}} الى {{travel_end_date}}"),
    (
        "l",
        "نفقات الرحلة: سيتكفل {{companion_name}} بمصاريف {{worker_name}} خلال فترة "
        "إقامتها في {{destination_country}}.",
    ),
    ("s", ""),
    ("h", "مع خالص الاحترام،"),
    ("b", "HR Manager"),
    ("b", "+971 505544143"),
]


# Unicode bidi isolates, so a left-to-right run inside right-to-left text stays one unit.
_LTR_ISOLATE = "⁦"  # LEFT-TO-RIGHT ISOLATE
_POP_ISOLATE = "⁩"  # POP DIRECTIONAL ISOLATE

# A run of Latin/digits and the punctuation that belongs inside it (phone "+", date comma,
# EID hyphens, "P.O. Box"). A leading "+" must be inside the isolate, not beside it — that is
# the character bidi was moving to the wrong end of the phone number.
_LTR_RUN_RE = re.compile(r"\+?[A-Za-z0-9][A-Za-z0-9 .,:+/()–-]*[A-Za-z0-9)]|\+?[A-Za-z0-9]")
_PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")
_DIGIT_RE = re.compile(r"[0-9]")

# U+200F RIGHT-TO-LEFT MARK. A paragraph whose whole content is left-to-right (the date line,
# the phone) otherwise takes its line placement from that content and lands against the left
# margin, breaking the page's right-to-left alignment. This anchors the paragraph direction
# without printing anything.
_RTL_MARK = "‏"


def _isolate_ltr_runs(text: str) -> str:
    """Fence each left-to-right run in an Arabic line so bidi cannot reorder it.

    Without this the Arabic page reorders its own static text: the signature phone
    "+971 505544143" printed as "505544143 971+", because bidi reads the "+" as trailing the
    number rather than leading it. Same class of problem as the substituted values, but these
    are authored into the template, so doc_utils never sees them.

    Placeholders are left alone: doc_utils isolates each substituted value itself, and wrapping
    the ``{{name}}`` token here would just nest the isolates.

    The line is also prefixed with a right-to-left mark, so a paragraph that holds nothing but a
    left-to-right value still sits against the right margin like the rest of the page.
    """
    def isolate(segment: str) -> str:
        # Only runs containing a number need fencing; a plain Latin word such as "HR Manager"
        # is already strongly left-to-right and renders correctly inside Arabic on its own.
        return _LTR_RUN_RE.sub(
            lambda m: f"{_LTR_ISOLATE}{m.group(0)}{_POP_ISOLATE}"
            if _DIGIT_RE.search(m.group(0)) else m.group(0),
            segment,
        )

    out, pos = [], 0
    for placeholder in _PLACEHOLDER_RE.finditer(text):
        out.append(isolate(text[pos:placeholder.start()]))
        out.append(placeholder.group(0))
        pos = placeholder.end()
    out.append(isolate(text[pos:]))
    return _RTL_MARK + "".join(out)


def _require_logo() -> Path:
    """The letterhead, extracted once from the retired PDF and committed alongside this script."""
    if not LOGO.exists():
        raise SystemExit(f"Missing letterhead asset: {LOGO.relative_to(BASE_DIR)}")
    return LOGO


def _set_bidi(paragraph, rtl: bool) -> None:
    """Mark the paragraph right-to-left.

    ``doc_utils.fill_document`` rebuilds a paragraph's runs when it substitutes a
    placeholder, so run-level ``w:rtl`` would be lost. Direction has to live on the
    paragraph (``w:bidi``) and the formatting on the paragraph *style*, which the rebuilt
    runs inherit.

    Alignment is deliberately NOT set. In a ``w:bidi`` paragraph ``w:jc`` is logical, not
    physical: ``w:val="right"`` means "end", which in right-to-left text is the *left* edge.
    Setting it to "right" therefore left-aligned the whole Arabic page — invisible on the long
    wrapped paragraphs, which fill the column either way, but plain on the headings and on any
    line holding only a date or the phone number. Leaving ``w:jc`` off falls back to "start",
    which is the right edge for an RTL paragraph and the left edge for an LTR one — correct for
    both pages, in Word and LibreOffice alike. Callers that genuinely want centring (the title,
    the footer) set ``paragraph.alignment`` themselves afterwards.
    """
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:bidi", "w:jc"):
        existing = p_pr.find(qn(tag))
        if existing is not None:
            p_pr.remove(existing)
    if rtl:
        p_pr.append(OxmlElement("w:bidi"))


def _style(doc, name: str, *, size: int, bold: bool, rtl: bool, font: str = BODY_FONT):
    """A paragraph style whose font the rebuilt runs will inherit."""
    from docx.enum.style import WD_STYLE_TYPE

    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    # Complex-script font, so Arabic renders in the same face instead of a fallback.
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi"):
        r_fonts.set(qn(attr), font)
    r_fonts.set(qn("w:cs"), ARABIC_FONT if rtl else font)
    if rtl:
        for tag in ("w:rtl", "w:cs"):
            element = OxmlElement(tag)
            r_pr.append(element)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(size * 2))
        r_pr.append(sz_cs)
    fmt = style.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(2)
    fmt.line_spacing = 1.0
    return style


def _add_block(doc, entries, *, rtl: bool, styles: dict) -> None:
    for kind, text in entries:
        if kind == "s":
            spacer = doc.add_paragraph(style=styles["body"])
            spacer.paragraph_format.space_after = Pt(6)
            _set_bidi(spacer, rtl)
            continue
        style_name = {"h": "heading", "b": "body", "l": "bullet"}[kind]
        paragraph = doc.add_paragraph(style=styles[style_name])
        _set_bidi(paragraph, rtl)
        # Only the Arabic page needs bidi isolation; on the English page it would be inert noise.
        body = _isolate_ltr_runs(text) if rtl else text
        if kind == "l":
            paragraph.paragraph_format.left_indent = Inches(0 if rtl else 0.35)
            paragraph.paragraph_format.right_indent = Inches(0.35 if rtl else 0)
            paragraph.add_run("- " + body)
        else:
            paragraph.add_run(body)
        if kind == "h":
            paragraph.runs[0].underline = True


def _add_letterhead(doc, title: str, *, rtl: bool, styles: dict) -> None:
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.add_run().add_picture(str(LOGO), width=Pt(LOGO_WIDTH_PT))
    heading = doc.add_paragraph(style=styles["title"])
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_bidi(heading, rtl)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(title)


def _build_footer(section, styles) -> None:
    """The attestation + address + disclaimer block, repeated on both pages."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for paragraph in list(footer.paragraphs):
        paragraph._p.getparent().remove(paragraph._p)
    for text, size, rtl in (
        (FOOTER_ATTESTATION, 9, False),
        (FOOTER_ADDRESS_AR, 9, True),
        (FOOTER_PHONE, 9, False),
        (FOOTER_DISCLAIMER, 6, False),
    ):
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if rtl:
            _set_bidi(paragraph, True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(_isolate_ltr_runs(text) if rtl else text)
        run.font.name = BODY_FONT
        run.font.size = Pt(size)
        r_pr = run._r.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:cs"), ARABIC_FONT if rtl else BODY_FONT)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(size * 2))
        r_pr.append(sz_cs)


def main() -> None:
    _require_logo()
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)

    styles = {
        "title": _style(doc, "NocTitle", size=TITLE_SIZE, bold=True, rtl=False, font=TITLE_FONT),
        "heading": _style(doc, "NocHeading", size=BODY_SIZE, bold=True, rtl=False),
        "body": _style(doc, "NocBody", size=BODY_SIZE, bold=False, rtl=False),
        "bullet": _style(doc, "NocBullet", size=BODY_SIZE, bold=False, rtl=False),
        "title_ar": _style(doc, "NocTitleAr", size=TITLE_SIZE, bold=True, rtl=True, font=TITLE_FONT),
        "heading_ar": _style(doc, "NocHeadingAr", size=BODY_SIZE, bold=True, rtl=True),
        "body_ar": _style(doc, "NocBodyAr", size=BODY_SIZE, bold=False, rtl=True),
        "bullet_ar": _style(doc, "NocBulletAr", size=BODY_SIZE, bold=False, rtl=True),
    }

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1.6)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.footer_distance = Inches(0.35)
    _build_footer(section, styles)

    latin = {k: styles[k] for k in ("title", "heading", "body", "bullet")}
    arabic = {
        "title": styles["title_ar"],
        "heading": styles["heading_ar"],
        "body": styles["body_ar"],
        "bullet": styles["bullet_ar"],
    }

    _add_letterhead(doc, ENGLISH_TITLE, rtl=False, styles=latin)
    _add_block(doc, ENGLISH_BODY, rtl=False, styles=latin)

    doc.add_section(WD_SECTION.NEW_PAGE)
    second = doc.sections[1]
    second.page_width = Inches(8.5)
    second.page_height = Inches(11)
    second.top_margin = Inches(0.5)
    second.bottom_margin = Inches(1.6)
    second.left_margin = Inches(1)
    second.right_margin = Inches(1)
    second.footer_distance = Inches(0.35)
    second.footer.is_linked_to_previous = True

    _add_letterhead(doc, ARABIC_TITLE, rtl=True, styles=arabic)
    _add_block(doc, ARABIC_BODY, rtl=True, styles=arabic)

    doc.save(OUTPUT)
    print(f"  letterhead -> {LOGO.relative_to(BASE_DIR)}")
    print(f"Done -> {OUTPUT.name}")


if __name__ == "__main__":
    main()
